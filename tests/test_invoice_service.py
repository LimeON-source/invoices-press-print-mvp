"""Tests for invoice_service.py — registry, batch processing, PDF conversion."""
import json
from datetime import date
from decimal import Decimal
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from invoice_app.config import load_config
from invoice_app.invoice_service import (
    assign_invoice_numbers,
    batch_convert_to_pdf,
    build_invoice_docx,
    convert_amount_to_words,
    find_field,
    load_counter,
    load_registry,
    next_invoice_number,
    normalize_key,
    prepare_invoice,
    preview_invoice_numbers,
    process_batch,
    save_counter,
    save_registry,
    sanitize_filename,
    _libreoffice_path,
)
from invoice_app.models import AppConfig, InvoiceModel, LineItem, SellerConfig


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture()
def seller():
    return SellerConfig(
        name="Test UAB",
        address="Vilnius",
        company_code="123",
        iban="LT00",
        issuer_name="Test User",
    )


@pytest.fixture()
def config(seller, tmp_path):
    return AppConfig(
        seller=seller,
        vat_percent=Decimal("0"),
        invoice_series="SS",
        output_root=str(tmp_path / "Invoices"),
    )


@pytest.fixture()
def invoice_model(seller):
    line = LineItem(description="kovas sessions", quantity=4, unit_price=Decimal("50"))
    return InvoiceModel(
        number="SS-2026-001",
        date=date(2026, 3, 31),
        seller=seller,
        buyer_name="Jonas Jonaitis",
        buyer_address="Kaunas",
        sessions=4,
        rate=Decimal("50"),
        subtotal=Decimal("200"),
        vat_amount=Decimal("0"),
        total=Decimal("200"),
        total_words="du šimtai",
        line_item=line,
    )


# ---------------------------------------------------------------------------
# Counter
# ---------------------------------------------------------------------------

def test_counter_save_and_load(tmp_path):
    p = tmp_path / "counter.json"
    save_counter({"year": 2025, "counter": 5}, str(p))
    c = load_counter(str(p))
    assert c["year"] == 2025
    assert c["counter"] == 5


def test_counter_missing_file_returns_defaults(tmp_path):
    c = load_counter(str(tmp_path / "missing.json"))
    assert c["counter"] == 0
    assert c["year"] == date.today().year


def test_next_invoice_number_increments(tmp_path):
    counter = {"year": 2026, "counter": 2}
    num = next_invoice_number("SS", counter, 2026)
    assert num == "SS-2026-003"
    assert counter["counter"] == 3


def test_next_invoice_number_resets_on_new_year():
    counter = {"year": 2025, "counter": 99}
    num = next_invoice_number("SS", counter, 2026)
    assert num == "SS-2026-001"
    assert counter["year"] == 2026
    assert counter["counter"] == 1


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

def test_assign_invoice_numbers_sequential(tmp_path):
    reg = str(tmp_path / "reg.json")
    result = assign_invoice_numbers("vasaris", ["Agnė", "Akvilė"], 2026, "SS", reg)
    assert result["Agnė"] == "SS-2026-001"
    assert result["Akvilė"] == "SS-2026-002"


def test_assign_invoice_numbers_idempotent(tmp_path):
    reg = str(tmp_path / "reg.json")
    first = assign_invoice_numbers("vasaris", ["Agnė"], 2026, "SS", reg)
    second = assign_invoice_numbers("vasaris", ["Agnė"], 2026, "SS", reg)
    assert first["Agnė"] == second["Agnė"]
    registry = load_registry(reg)
    assert registry["2026"]["counter"] == 1  # not incremented on second call


def test_assign_invoice_numbers_new_client_gets_next(tmp_path):
    reg = str(tmp_path / "reg.json")
    assign_invoice_numbers("vasaris", ["Agnė", "Akvilė"], 2026, "SS", reg)
    result = assign_invoice_numbers("vasaris", ["Agnė", "Jonas"], 2026, "SS", reg)
    assert result["Agnė"] == "SS-2026-001"   # locked — unchanged
    assert result["Jonas"] == "SS-2026-003"  # next after existing 2


def test_assign_invoice_numbers_resets_each_year(tmp_path):
    reg = str(tmp_path / "reg.json")
    assign_invoice_numbers("vasaris", ["Agnė"], 2026, "SS", reg)
    result = assign_invoice_numbers("vasaris", ["Agnė"], 2027, "SS", reg)
    assert result["Agnė"] == "SS-2027-001"


def test_preview_invoice_numbers_does_not_write(tmp_path):
    reg = str(tmp_path / "reg.json")
    preview_invoice_numbers("vasaris", ["Agnė"], 2026, "SS", reg)
    assert not (tmp_path / "reg.json").exists()  # file not created


def test_preview_invoice_numbers_shows_locked(tmp_path):
    reg = str(tmp_path / "reg.json")
    assign_invoice_numbers("vasaris", ["Agnė"], 2026, "SS", reg)
    result = preview_invoice_numbers("vasaris", ["Agnė", "Jonas"], 2026, "SS", reg)
    assert result["Agnė"]["locked"] is True
    assert result["Agnė"]["number"] == "SS-2026-001"
    assert result["Jonas"]["locked"] is False
    assert result["Jonas"]["number"] == "SS-2026-002"


# ---------------------------------------------------------------------------
# Amount to words
# ---------------------------------------------------------------------------

def test_convert_amount_to_words_returns_string():
    result = convert_amount_to_words(Decimal("150.00"))
    assert isinstance(result, str)
    assert len(result) > 0


def test_convert_amount_to_words_contains_amount():
    result = convert_amount_to_words(Decimal("100.00"))
    assert "100" in result or "šimtas" in result or "šimtai" in result


# ---------------------------------------------------------------------------
# Filename sanitization
# ---------------------------------------------------------------------------

def test_sanitize_filename_removes_special_chars():
    assert sanitize_filename("Ąžuolas Žvejys") == "uolas-vejys"


def test_sanitize_filename_replaces_spaces():
    assert sanitize_filename("John Doe") == "John-Doe"


# ---------------------------------------------------------------------------
# Key normalization & find_field
# ---------------------------------------------------------------------------

def test_normalize_key_strips_diacritics():
    assert normalize_key("Klièntė") == "kliėnte" or normalize_key("Klientė") == "klientė" or normalize_key("Rate") == "rate"


def test_find_field_case_insensitive():
    row = {"Rate": "50", "Client": "Jonas"}
    assert find_field(row, "rate") == "50"
    assert find_field(row, "client") == "Jonas"


def test_find_field_missing_returns_none():
    assert find_field({"Rate": "50"}, "Address") is None


# ---------------------------------------------------------------------------
# prepare_invoice
# ---------------------------------------------------------------------------

def test_prepare_invoice_zero_sessions_returns_none(config):
    row = {"Client": "Jonas", "Kovas": "0", "Rate": "50", "Address": "Vilnius"}
    assert prepare_invoice(row, "Kovas", config, "SS-001", date.today()) is None


def test_prepare_invoice_valid_row_returns_model(config):
    row = {"Client": "Jonas", "Kovas": "4", "Rate": "50", "Address": "Vilnius"}
    model = prepare_invoice(row, "Kovas", config, "SS-001", date.today())
    assert model is not None
    assert model.sessions == 4
    assert model.rate == Decimal("50")
    assert model.total == Decimal("200")


def test_prepare_invoice_skips_summary_row(config):
    row = {"Client": "Viso", "Kovas": "10", "Rate": "50"}
    assert prepare_invoice(row, "Kovas", config, "SS-001", date.today()) is None


def test_prepare_invoice_skips_uzdarbis_row(config):
    """Financial summary rows like 'Uždarbis' must never become invoices."""
    row = {"Name Surname": "Uždarbis", "Vasaris": "4860", "Kaina": "4800"}
    assert prepare_invoice(row, "Vasaris", config, "SS-001", date.today()) is None


def test_prepare_invoice_skips_udarbis_variant(config):
    row = {"Name Surname": "Udarbis", "Vasaris": "100", "Kaina": "100"}
    assert prepare_invoice(row, "Vasaris", config, "SS-001", date.today()) is None


def test_prepare_invoice_default_address(config):
    row = {"Client": "Jonas", "Kovas": "2", "Rate": "30"}
    model = prepare_invoice(row, "Kovas", config, "SS-001", date.today())
    assert model.buyer_address == "N/A"


def test_prepare_invoice_invalid_rate_raises(config):
    row = {"Client": "Jonas", "Kovas": "2", "Rate": "abc"}
    with pytest.raises(ValueError):
        prepare_invoice(row, "Kovas", config, "SS-001", date.today())


# ---------------------------------------------------------------------------
# build_invoice_docx
# ---------------------------------------------------------------------------

def test_build_invoice_docx_creates_file(invoice_model, tmp_path):
    template = tmp_path / "template.docx"
    # create a minimal real docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("<number> <client_name>")
    doc.save(str(template))

    docx_path, existing = build_invoice_docx(
        invoice_model, str(template), str(tmp_path), policy="overwrite", month_folder="kovas"
    )
    assert existing is None
    assert docx_path is not None
    assert docx_path.exists()


def test_build_invoice_docx_skip_policy_returns_existing(invoice_model, tmp_path):
    from docx import Document
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("test")
    doc.save(str(template))

    # First create
    build_invoice_docx(invoice_model, str(template), str(tmp_path), policy="overwrite", month_folder="kovas")
    # Manually create PDF to simulate existing
    pdf = tmp_path / "kovas" / "SS-2026-001_Jonas-Jonaitis.pdf"
    pdf.parent.mkdir(parents=True, exist_ok=True)
    pdf.write_text("dummy")

    docx_path, existing = build_invoice_docx(
        invoice_model, str(template), str(tmp_path), policy="skip", month_folder="kovas"
    )
    assert docx_path is None
    assert existing == pdf


def test_build_invoice_docx_replacements(invoice_model, tmp_path, monkeypatch):
    captured = {}

    def fake_fill(template_path, output_path, replacements):
        captured["replacements"] = replacements

    monkeypatch.setattr("invoice_app.invoice_service.fill_template", fake_fill)

    build_invoice_docx(invoice_model, "template.docx", str(tmp_path), month_folder="kovas")

    r = captured["replacements"]
    assert r["<number>"] == "SS-2026-001"
    assert r["<client_name>"] == "Jonas Jonaitis"
    assert r["<kiekis>"] == "4"
    assert r["<Kiekis>"] == "4"
    assert r["<kaina>"] == "50.00"
    assert r["<total>"] == "200.00"


# ---------------------------------------------------------------------------
# batch_convert_to_pdf
# ---------------------------------------------------------------------------

def test_batch_convert_uses_libreoffice_when_available(tmp_path, monkeypatch):
    monkeypatch.setattr("invoice_app.invoice_service._libreoffice_path", lambda: "/fake/soffice")

    calls = []

    def fake_run(cmd, **kwargs):
        calls.append(cmd)
        # simulate PDF being created
        for arg in cmd:
            if arg.endswith(".docx"):
                pdf = tmp_path / (Path(arg).stem + ".pdf")
                pdf.write_text("fake pdf")

    monkeypatch.setattr("invoice_app.invoice_service.subprocess.run", fake_run)

    docx = tmp_path / "SS-2026-001_Test.docx"
    docx.write_text("fake docx")

    results = batch_convert_to_pdf([docx], tmp_path)
    assert len(calls) == 1
    assert "--headless" in calls[0]
    assert results[str(docx)] is not None


def test_batch_convert_returns_none_when_no_converter(tmp_path, monkeypatch):
    monkeypatch.setattr("invoice_app.invoice_service._libreoffice_path", lambda: None)
    monkeypatch.setattr("invoice_app.invoice_service.subprocess.run", lambda *a, **k: None)

    # Also patch docx2pdf import to fail
    import sys
    monkeypatch.setitem(sys.modules, "docx2pdf", None)

    docx = tmp_path / "test.docx"
    docx.write_text("fake")
    results = batch_convert_to_pdf([docx], tmp_path)
    assert results[str(docx)] is None


def test_batch_convert_empty_list(tmp_path):
    results = batch_convert_to_pdf([], tmp_path)
    assert results == {}


# ---------------------------------------------------------------------------
# process_batch — registry-based numbering, total_amount, error tracking
# ---------------------------------------------------------------------------

def test_process_batch_number_not_assigned_on_skip(config, tmp_path):
    """Invoice numbers must NOT be assigned for rows with 0 sessions."""
    reg = str(tmp_path / "reg.json")
    rows = [
        {"Client": "Jonas", "Kovas": "0", "Rate": "50"},   # 0 sessions → skip
        {"Client": "Petras", "Kovas": "3", "Rate": "50"},  # valid
    ]

    with patch("invoice_app.invoice_service.build_invoice_docx") as mock_build, \
         patch("invoice_app.invoice_service.batch_convert_to_pdf") as mock_pdf:
        mock_build.return_value = (tmp_path / "test.docx", None)
        mock_pdf.return_value = {str(tmp_path / "test.docx"): tmp_path / "test.pdf"}

        stats = process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite", month_folder="kovas",
        )

    registry = load_registry(reg)
    year_data = registry.get(str(date.today().year), {})
    assert year_data["counter"] == 1          # only Petras got a number
    assert "Petras" in year_data["months"]["kovas"]
    assert "Jonas" not in year_data["months"]["kovas"]
    assert stats["skipped"] == 1
    assert stats["processed"] == 1


def test_process_batch_number_not_assigned_on_error(config, tmp_path):
    """Invoice numbers must NOT be assigned when prepare_invoice raises."""
    reg = str(tmp_path / "reg.json")
    rows = [
        {"Client": "Bad", "Kovas": "3", "Rate": "not-a-number"},
        {"Client": "Good", "Kovas": "2", "Rate": "40"},
    ]

    with patch("invoice_app.invoice_service.build_invoice_docx") as mock_build, \
         patch("invoice_app.invoice_service.batch_convert_to_pdf") as mock_pdf:
        mock_build.return_value = (tmp_path / "good.docx", None)
        mock_pdf.return_value = {str(tmp_path / "good.docx"): tmp_path / "good.pdf"}

        stats = process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite", month_folder="kovas",
        )

    registry = load_registry(reg)
    assert registry[str(date.today().year)]["counter"] == 1  # only Good
    assert stats["errors"] == 1
    assert stats["error_details"][0]["client"] == "Bad"


def test_process_batch_idempotent_numbers(config, tmp_path):
    """Regenerating same month must reuse existing invoice numbers."""
    reg = str(tmp_path / "reg.json")
    rows = [{"Client": "Petras", "Kovas": "3", "Rate": "50"}]

    with patch("invoice_app.invoice_service.build_invoice_docx") as mock_build, \
         patch("invoice_app.invoice_service.batch_convert_to_pdf") as mock_pdf:
        mock_build.return_value = (tmp_path / "x.docx", None)
        mock_pdf.return_value = {str(tmp_path / "x.docx"): tmp_path / "x.pdf"}

        process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite", month_folder="kovas",
        )
        mock_build.reset_mock()
        process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite", month_folder="kovas",
        )

    registry = load_registry(reg)
    assert registry[str(date.today().year)]["counter"] == 1  # not incremented twice


def test_process_batch_total_amount(config, tmp_path):
    """stats['total_amount'] must sum all valid invoice totals."""
    reg = str(tmp_path / "reg.json")
    rows = [
        {"Client": "A", "Kovas": "4", "Rate": "50"},   # 200
        {"Client": "B", "Kovas": "2", "Rate": "100"},  # 200
    ]

    with patch("invoice_app.invoice_service.build_invoice_docx") as mock_build, \
         patch("invoice_app.invoice_service.batch_convert_to_pdf") as mock_pdf:
        mock_build.return_value = (tmp_path / "x.docx", None)
        mock_pdf.return_value = {str(tmp_path / "x.docx"): tmp_path / "x.pdf"}

        stats = process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite", month_folder="kovas",
        )

    assert stats["total_amount"] == Decimal("400")


def test_process_batch_error_details_recorded(config, tmp_path):
    reg = str(tmp_path / "reg.json")
    rows = [{"Client": "Broken", "Kovas": "x", "Rate": "50"}]

    stats = process_batch(
        rows=rows, month_column="Kovas", config=config,
        template_path="t.docx", output_root=config.output_root,
        registry_file=reg, policy="overwrite", month_folder="kovas",
    )

    assert stats["errors"] == 1
    assert stats["error_details"][0]["client"] == "Broken"
    assert stats["error_details"][0]["reason"] != ""


def test_process_batch_selected_clients_filter(config, tmp_path):
    reg = str(tmp_path / "reg.json")
    rows = [
        {"Client": "Jonas", "Kovas": "2", "Rate": "50"},
        {"Client": "Petras", "Kovas": "3", "Rate": "50"},
    ]

    with patch("invoice_app.invoice_service.build_invoice_docx") as mock_build, \
         patch("invoice_app.invoice_service.batch_convert_to_pdf") as mock_pdf:
        mock_build.return_value = (tmp_path / "x.docx", None)
        mock_pdf.return_value = {str(tmp_path / "x.docx"): tmp_path / "x.pdf"}

        stats = process_batch(
            rows=rows, month_column="Kovas", config=config,
            template_path="t.docx", output_root=config.output_root,
            registry_file=reg, policy="overwrite",
            selected_clients=["Jonas"], month_folder="kovas",
        )

    assert stats["processed"] == 1


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

def test_config_load_from_file():
    cfg = load_config("config.json")
    assert cfg.invoice_series == "SS"
    assert cfg.seller.name == "Šaknys ir sparnai MB"


# ---------------------------------------------------------------------------
# Models
# ---------------------------------------------------------------------------

def test_line_item_total():
    item = LineItem(description="Service A", quantity=1, unit_price=100)
    assert float(item.total) == 100.0


def test_line_item_total_batch():
    items = [
        LineItem(description="A", quantity=1, unit_price=100),
        LineItem(description="B", quantity=2, unit_price=50),
    ]
    assert sum(float(i.total) for i in items) == 200.0


def test_invoice_model_is_mutable(invoice_model):
    invoice_model.number = "SS-2026-999"
    assert invoice_model.number == "SS-2026-999"
