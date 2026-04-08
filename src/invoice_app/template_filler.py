from pathlib import Path
from typing import Dict

from docx import Document


def replace_text_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    text = paragraph.text
    for key, value in replacements.items():
        if key in text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)


def replace_text_in_table(table, replacements: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def fill_template(template_path: str, output_path: str, replacements: Dict[str, str]) -> None:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        replace_text_in_table(table, replacements)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
